#!/usr/bin/env python3
"""Fill placeholder values in a Word document template."""

import json
import sys
import os
import re


def ensure_dependencies():
    try:
        import docx
    except ImportError:
        os.system("pip install python-docx")


def find_placeholders(doc):
    """Find all {{placeholder}} patterns in the document."""
    placeholders = set()
    pattern = re.compile(r'\{\{(\w+)\}\}')

    for para in doc.paragraphs:
        # Check full paragraph text (handles split runs)
        matches = pattern.findall(para.text)
        placeholders.update(matches)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = pattern.findall(para.text)
                    placeholders.update(matches)

    for section in doc.sections:
        if section.header:
            for para in section.header.paragraphs:
                matches = pattern.findall(para.text)
                placeholders.update(matches)
        if section.footer:
            for para in section.footer.paragraphs:
                matches = pattern.findall(para.text)
                placeholders.update(matches)

    return sorted(list(placeholders))


def replace_placeholder_in_runs(paragraph, placeholder, value):
    """Replace a placeholder that may be split across multiple runs."""
    full_text = paragraph.text
    target = '{{' + placeholder + '}}'

    if target not in full_text:
        return False

    # Try simple run-level replacement first
    for run in paragraph.runs:
        if target in run.text:
            run.text = run.text.replace(target, str(value))
            return True

    # Handle placeholder split across runs
    runs_text = [(run, run.text) for run in paragraph.runs]
    combined = ''.join(t for _, t in runs_text)

    if target in combined:
        new_combined = combined.replace(target, str(value))
        # Clear all runs and put result in first
        for i, (run, _) in enumerate(runs_text):
            if i == 0:
                run.text = new_combined
            else:
                run.text = ''
        return True

    return False


def fill_template(template_path, values, output_path):
    """Fill all placeholders in a template document."""
    ensure_dependencies()
    from docx import Document

    doc = Document(template_path)

    # Find all placeholders first
    all_placeholders = find_placeholders(doc)

    # Track what was filled and what was missing
    filled = []
    missing_values = []
    unused_values = []

    for ph in all_placeholders:
        if ph in values:
            filled.append(ph)
        else:
            missing_values.append(ph)

    for key in values:
        if key not in all_placeholders:
            unused_values.append(key)

    # Do replacements
    for placeholder, value in values.items():
        for para in doc.paragraphs:
            replace_placeholder_in_runs(para, placeholder, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_placeholder_in_runs(para, placeholder, value)

        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    replace_placeholder_in_runs(para, placeholder, value)
            if section.footer:
                for para in section.footer.paragraphs:
                    replace_placeholder_in_runs(para, placeholder, value)

    os.makedirs(os.path.dirname(output_path) or '/tmp', exist_ok=True)
    doc.save(output_path)

    result = {
        "status": "success",
        "output_path": output_path,
        "placeholders_found": all_placeholders,
        "filled": filled,
    }
    if missing_values:
        result["missing_values"] = missing_values
        result["warning"] = f"No values provided for: {', '.join(missing_values)}"
    if unused_values:
        result["unused_values"] = unused_values

    return result


def main():
    if len(sys.argv) < 2:
        print(json.dumps({
            "error": "Usage: template_fill.py <template.docx> --values '{\"key\": \"value\"}' --output <output.docx>"
        }))
        sys.exit(1)

    ensure_dependencies()

    template_path = sys.argv[1]

    if not os.path.exists(template_path):
        print(json.dumps({"error": f"Template not found: {template_path}"}))
        sys.exit(1)

    # Scan-only mode
    if "--scan" in sys.argv:
        from docx import Document
        doc = Document(template_path)
        placeholders = find_placeholders(doc)
        print(json.dumps({"placeholders": placeholders, "count": len(placeholders)}, indent=2))
        return

    values = {}
    output_path = template_path.replace('.docx', '_filled.docx')

    if "--values" in sys.argv:
        idx = sys.argv.index("--values")
        if idx + 1 < len(sys.argv):
            values_arg = sys.argv[idx + 1]
            if os.path.isfile(values_arg):
                with open(values_arg, 'r') as f:
                    values = json.load(f)
            else:
                values = json.loads(values_arg)

    if "--output" in sys.argv:
        idx = sys.argv.index("--output")
        if idx + 1 < len(sys.argv):
            output_path = sys.argv[idx + 1]

    result = fill_template(template_path, values, output_path)
    print(json.dumps(result, indent=2))


if __name__ == '__main__':
    main()
