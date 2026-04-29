#!/usr/bin/env python3
"""Read and extract content from Word documents (.docx)."""

import json
import sys
import os
import base64


def ensure_dependencies():
    try:
        import docx
    except ImportError:
        os.system("pip install python-docx")


def extract_text(doc_path):
    """Extract plain text from all paragraphs."""
    from docx import Document
    doc = Document(doc_path)
    lines = []
    for para in doc.paragraphs:
        lines.append(para.text)
    return '\n'.join(lines)


def extract_structured(doc_path):
    """Extract structured content with formatting info."""
    from docx import Document
    doc = Document(doc_path)

    content = []
    for para in doc.paragraphs:
        runs = []
        for run in para.runs:
            run_info = {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
            }
            if run.font.name:
                run_info["font_name"] = run.font.name
            if run.font.size:
                run_info["font_size"] = run.font.size.pt
            if run.font.color and run.font.color.rgb:
                run_info["color"] = str(run.font.color.rgb)
            runs.append(run_info)

        para_info = {
            "text": para.text,
            "style": para.style.name if para.style else None,
            "alignment": str(para.alignment) if para.alignment else None,
            "runs": runs,
        }
        content.append(para_info)

    # Extract headers and footers
    headers_footers = []
    for section in doc.sections:
        section_info = {}
        if section.header and section.header.paragraphs:
            section_info["header"] = [p.text for p in section.header.paragraphs if p.text]
        if section.footer and section.footer.paragraphs:
            section_info["footer"] = [p.text for p in section.footer.paragraphs if p.text]
        if section_info:
            headers_footers.append(section_info)

    return {
        "paragraphs": content,
        "headers_footers": headers_footers,
    }


def extract_tables(doc_path):
    """Extract all tables as arrays."""
    from docx import Document
    doc = Document(doc_path)

    tables = []
    for i, table in enumerate(doc.tables):
        table_data = {
            "table_index": i,
            "rows": []
        }
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data["rows"].append(row_data)
        tables.append(table_data)

    return {"tables": tables, "count": len(tables)}


def extract_metadata(doc_path):
    """Extract document metadata/properties."""
    from docx import Document
    doc = Document(doc_path)
    props = doc.core_properties

    meta = {
        "title": props.title,
        "author": props.author,
        "subject": props.subject,
        "keywords": props.keywords,
        "created": str(props.created) if props.created else None,
        "modified": str(props.modified) if props.modified else None,
        "last_modified_by": props.last_modified_by,
        "revision": props.revision,
        "category": props.category,
        "comments": props.comments,
    }

    # Stats
    para_count = len(doc.paragraphs)
    word_count = sum(len(p.text.split()) for p in doc.paragraphs)
    char_count = sum(len(p.text) for p in doc.paragraphs)
    table_count = len(doc.tables)
    section_count = len(doc.sections)

    meta["statistics"] = {
        "paragraph_count": para_count,
        "word_count": word_count,
        "character_count": char_count,
        "table_count": table_count,
        "section_count": section_count,
    }

    # Styles used
    styles_used = set()
    for p in doc.paragraphs:
        if p.style:
            styles_used.add(p.style.name)
    meta["styles_used"] = sorted(list(styles_used))

    return meta


def extract_images(doc_path, output_dir="/tmp/docx_images"):
    """Extract embedded images to output directory."""
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    os.makedirs(output_dir, exist_ok=True)
    doc = Document(doc_path)

    image_paths = []
    image_count = 0

    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
            image_data = rel.target_part.blob
            ext = os.path.splitext(rel.target_part.partname)[1] or '.png'
            filename = f"image_{image_count}{ext}"
            filepath = os.path.join(output_dir, filename)
            with open(filepath, 'wb') as f:
                f.write(image_data)
            image_paths.append(filepath)

    return {"images": image_paths, "count": image_count}


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: read_docx.py <file.docx> [--mode text|structured|tables|metadata|images]"}))
        sys.exit(1)

    ensure_dependencies()

    doc_path = sys.argv[1]
    mode = "text"

    if "--mode" in sys.argv:
        idx = sys.argv.index("--mode")
        if idx + 1 < len(sys.argv):
            mode = sys.argv[idx + 1]

    if not os.path.exists(doc_path):
        print(json.dumps({"error": f"File not found: {doc_path}"}))
        sys.exit(1)

    try:
        if mode == "text":
            result = {"text": extract_text(doc_path)}
        elif mode == "structured":
            result = extract_structured(doc_path)
        elif mode == "tables":
            result = extract_tables(doc_path)
        elif mode == "metadata":
            result = extract_metadata(doc_path)
        elif mode == "images":
            output_dir = "/tmp/docx_images"
            if "--output" in sys.argv:
                idx = sys.argv.index("--output")
                if idx + 1 < len(sys.argv):
                    output_dir = sys.argv[idx + 1]
            result = extract_images(doc_path, output_dir)
        else:
            result = {"error": f"Unknown mode: {mode}. Use: text, structured, tables, metadata, images"}

        print(json.dumps(result, indent=2, default=str))
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)


if __name__ == '__main__':
    main()
