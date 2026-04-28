#!/usr/bin/env python3
"""Convert Word documents to/from various formats."""

import json
import sys
import os
import subprocess
import shutil


def ensure_dependencies():
    try:
        import docx
    except ImportError:
        os.system("pip install python-docx")


def docx_to_text(input_path, output_path):
    """Convert .docx to plain text."""
    ensure_dependencies()
    from docx import Document

    doc = Document(input_path)
    lines = [para.text for para in doc.paragraphs]
    text = '\n'.join(lines)

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)

    return {"status": "success", "output_path": output_path, "word_count": len(text.split())}


def docx_to_markdown(input_path, output_path):
    """Convert .docx to Markdown."""
    ensure_dependencies()
    from docx import Document

    doc = Document(input_path)
    md_lines = []

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''

        if style_name.startswith('Heading'):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            md_lines.append(f"{'#' * level} {para.text}")
            md_lines.append('')

        elif style_name == 'List Bullet':
            md_lines.append(f"- {para.text}")

        elif style_name == 'List Number':
            md_lines.append(f"1. {para.text}")

        elif para.text.strip():
            # Process inline formatting
            text_parts = []
            for run in para.runs:
                t = run.text
                if run.bold and run.italic:
                    t = f"***{t}***"
                elif run.bold:
                    t = f"**{t}**"
                elif run.italic:
                    t = f"*{t}*"
                text_parts.append(t)

            md_lines.append(''.join(text_parts) if text_parts else para.text)
            md_lines.append('')
        else:
            md_lines.append('')

    # Tables
    for table in doc.tables:
        md_lines.append('')
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('|', '\\|') for cell in row.cells]
            md_lines.append('| ' + ' | '.join(cells) + ' |')
            if i == 0:
                md_lines.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
        md_lines.append('')

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(md_lines))

    return {"status": "success", "output_path": output_path}


def docx_to_html(input_path, output_path):
    """Convert .docx to HTML."""
    ensure_dependencies()
    from docx import Document

    doc = Document(input_path)
    html_parts = ['<!DOCTYPE html>', '<html>', '<head>',
                  '<meta charset="utf-8">',
                  f'<title>{doc.core_properties.title or "Document"}</title>',
                  '<style>',
                  'body { font-family: Calibri, Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }',
                  'table { border-collapse: collapse; width: 100%; margin: 1em 0; }',
                  'th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }',
                  'th { background-color: #f2f2f2; font-weight: bold; }',
                  '</style>',
                  '</head>', '<body>']

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''

        if style_name.startswith('Heading'):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            level = min(level, 6)
            html_parts.append(f'<h{level}>{_escape_html(para.text)}</h{level}>')

        elif style_name == 'List Bullet':
            html_parts.append(f'<ul><li>{_escape_html(para.text)}</li></ul>')

        elif style_name == 'List Number':
            html_parts.append(f'<ol><li>{_escape_html(para.text)}</li></ol>')

        elif para.text.strip():
            inline_html = []
            for run in para.runs:
                t = _escape_html(run.text)
                if run.bold:
                    t = f'<strong>{t}</strong>'
                if run.italic:
                    t = f'<em>{t}</em>'
                if run.underline:
                    t = f'<u>{t}</u>'
                inline_html.append(t)
            html_parts.append(f'<p>{"".join(inline_html)}</p>')

    for table in doc.tables:
        html_parts.append('<table>')
        for i, row in enumerate(table.rows):
            html_parts.append('<tr>')
            tag = 'th' if i == 0 else 'td'
            for cell in row.cells:
                html_parts.append(f'<{tag}>{_escape_html(cell.text)}</{tag}>')
            html_parts.append('</tr>')
        html_parts.append('</table>')

    html_parts.extend(['</body>', '</html>'])

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(html_parts))

    return {"status": "success", "output_path": output_path}


def docx_to_pdf(input_path, output_path):
    """Convert .docx to PDF using LibreOffice."""
    lo_path = shutil.which('libreoffice') or shutil.which('soffice')
    if not lo_path:
        return {
            "status": "error",
            "message": "LibreOffice not found. Install it for PDF conversion: brew install --cask libreoffice (macOS) or apt install libreoffice (Linux)"
        }

    output_dir = os.path.dirname(output_path) or '/tmp'
    try:
        result = subprocess.run(
            [lo_path, '--headless', '--convert-to', 'pdf', '--outdir', output_dir, input_path],
            capture_output=True, text=True, timeout=120
        )
        if result.returncode != 0:
            return {"status": "error", "message": result.stderr}

        # LibreOffice names the output based on input filename
        expected_name = os.path.splitext(os.path.basename(input_path))[0] + '.pdf'
        expected_path = os.path.join(output_dir, expected_name)

        if expected_path != output_path and os.path.exists(expected_path):
            shutil.move(expected_path, output_path)

        return {"status": "success", "output_path": output_path}
    except subprocess.TimeoutExpired:
        return {"status": "error", "message": "PDF conversion timed out"}
    except Exception as e:
        return {"status": "error", "message": str(e)}


def markdown_to_docx(input_path, output_path):
    """Convert Markdown to .docx."""
    ensure_dependencies()
    from docx import Document
    import re

    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    lines = content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            doc.add_heading(heading_match.group(2), level=level)
            i += 1
            continue

        # Bullet list
        bullet_match = re.match(r'^[-*+]\s+(.*)', line)
        if bullet_match:
            doc.add_paragraph(bullet_match.group(1), style='List Bullet')
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^\d+\.\s+(.*)', line)
        if num_match:
            doc.add_paragraph(num_match.group(1), style='List Number')
            i += 1
            continue

        # Table
        if '|' in line and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|', lines[i + 1]):
            headers = [c.strip() for c in line.strip('|').split('|')]
            i += 2  # skip header separator
            rows = []
            while i < len(lines) and '|' in lines[i]:
                row = [c.strip() for c in lines[i].strip('|').split('|')]
                rows.append(row)
                i += 1
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            for j, h in enumerate(headers):
                table.rows[0].cells[j].text = h
            for row in rows:
                row_cells = table.add_row().cells
                for j, val in enumerate(row):
                    if j < len(headers):
                        row_cells[j].text = val
            continue

        # Regular paragraph
        if line.strip():
            p = doc.add_paragraph()
            # Parse inline formatting
            parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`)', line)
            for part in parts:
                if part.startswith('***') and part.endswith('***'):
                    run = p.add_run(part[3:-3])
                    run.bold = True
                    run.italic = True
                elif part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*'):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                elif part:
                    p.add_run(part)

        i += 1

    doc.save(output_path)
    return {"status": "success", "output_path": output_path}


def text_to_docx(input_path, output_path):
    """Convert plain text to .docx."""
    ensure_dependencies()
    from docx import Document

    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    for line in content.split('\n'):
        doc.add_paragraph(line)

    doc.save(output_path)
    return {"status": "success", "output_path": output_path}


def html_to_docx(input_path, output_path):
    """Convert HTML to .docx (basic conversion)."""
    ensure_dependencies()
    from docx import Document
    import re

    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Strip tags but preserve structure
    doc = Document()

    # Extract headings
    for match in re.finditer(r'<h([1-6])[^>]*>(.*?)</h\1>', content, re.DOTALL):
        level = int(match.group(1))
        text = re.sub(r'<[^>]+>', '', match.group(2)).strip()
        if text:
            doc.add_heading(text, level=level)

    # Extract paragraphs
    for match in re.finditer(r'<p[^>]*>(.*?)</p>', content, re.DOTALL):
        text = re.sub(r'<[^>]+>', '', match.group(1)).strip()
        if text:
            doc.add_paragraph(text)

    # Extract tables
    for table_match in re.finditer(r'<table[^>]*>(.*?)</table>', content, re.DOTALL):
        rows_html = re.findall(r'<tr[^>]*>(.*?)</tr>', table_match.group(1), re.DOTALL)
        if rows_html:
            first_row_cells = re.findall(r'<t[hd][^>]*>(.*?)</t[hd]>', rows_html[0], re.DOTALL)
            num_cols = len(first_row_cells)
            if num_cols > 0:
                table = doc.add_table(rows=0, cols=num_cols)
                table.style = 'Table Grid'
                for row_html in rows_html:
                    cells = re.findall(r'<t[hd][^>]*>(.*?)</t[hd]>', row_html, re.DOTALL)
                    row = table.add_row()
                    for i, cell_html in enumerate(cells):
                        if i < num_cols:
                            row.cells[i].text = re.sub(r'<[^>]+>', '', cell_html).strip()

    doc.save(output_path)
    return {"status": "success", "output_path": output_path}


def _escape_html(text):
    """Escape HTML special characters."""
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;')


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: convert_docx.py <input_file> --to <format> [--output <output_path>]"}))
        sys.exit(1)

    input_path = sys.argv[1]
    to_format = None
    output_path = None

    if "--to" in sys.argv:
        idx = sys.argv.index("--to")
        if idx + 1 < len(sys.argv):
            to_format = sys.argv[idx + 1].lower()

    if "--output" in sys.argv:
        idx = sys.argv.index("--output")
        if idx + 1 < len(sys.argv):
            output_path = sys.argv[idx + 1]

    if not to_format:
        print(json.dumps({"error": "Specify target format with --to (pdf, txt, html, md, docx)"}))
        sys.exit(1)

    if not os.path.exists(input_path):
        print(json.dumps({"error": f"File not found: {input_path}"}))
        sys.exit(1)

    input_ext = os.path.splitext(input_path)[1].lower()
    base_name = os.path.splitext(os.path.basename(input_path))[0]

    if not output_path:
        ext_map = {'pdf': '.pdf', 'txt': '.txt', 'html': '.html', 'md': '.md', 'docx': '.docx'}
        output_path = f"/tmp/{base_name}{ext_map.get(to_format, '.' + to_format)}"

    converters = {
        ('.docx', 'txt'): docx_to_text,
        ('.docx', 'md'): docx_to_markdown,
        ('.docx', 'html'): docx_to_html,
        ('.docx', 'pdf'): docx_to_pdf,
        ('.md', 'docx'): markdown_to_docx,
        ('.txt', 'docx'): text_to_docx,
        ('.html', 'docx'): html_to_docx,
        ('.htm', 'docx'): html_to_docx,
    }

    converter = converters.get((input_ext, to_format))
    if not converter:
        print(json.dumps({"error": f"Unsupported conversion: {input_ext} → {to_format}"}))
        sys.exit(1)

    result = converter(input_path, output_path)
    print(json.dumps(result, indent=2))


if __name__ == '__main__':
    main()
