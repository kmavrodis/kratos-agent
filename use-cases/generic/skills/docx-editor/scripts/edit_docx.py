#!/usr/bin/env python3
"""Edit existing Word documents (.docx) with various operations."""

import json
import sys
import os
import re
import copy


def ensure_dependencies():
    try:
        import docx
    except ImportError:
        os.system("pip install python-docx")


def find_replace(doc, find_text, replace_text, match_case=True):
    """Find and replace text across paragraphs and tables."""
    count = 0

    def replace_in_paragraph(paragraph):
        nonlocal count
        full_text = paragraph.text
        if not match_case:
            if find_text.lower() not in full_text.lower():
                return
        else:
            if find_text not in full_text:
                return

        # Rebuild runs to handle text split across runs
        for run in paragraph.runs:
            if match_case:
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_text)
                    count += 1
            else:
                pattern = re.compile(re.escape(find_text), re.IGNORECASE)
                if pattern.search(run.text):
                    run.text = pattern.sub(replace_text, run.text)
                    count += 1

        # Handle text split across runs
        if count == 0 and find_text in full_text:
            # Concatenate all runs, do replace, put result in first run
            combined = ''.join(run.text for run in paragraph.runs)
            if match_case:
                new_text = combined.replace(find_text, replace_text)
                occurrences = combined.count(find_text)
            else:
                pattern = re.compile(re.escape(find_text), re.IGNORECASE)
                occurrences = len(pattern.findall(combined))
                new_text = pattern.sub(replace_text, combined)

            if occurrences > 0:
                for i, run in enumerate(paragraph.runs):
                    if i == 0:
                        run.text = new_text
                    else:
                        run.text = ''
                count += occurrences

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

    # Headers and footers
    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                replace_in_paragraph(paragraph)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                replace_in_paragraph(paragraph)

    return count


def find_replace_regex(doc, pattern, replace_text):
    """Find and replace using regex patterns."""
    count = 0
    regex = re.compile(pattern)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            matches = regex.findall(run.text)
            if matches:
                run.text = regex.sub(replace_text, run.text)
                count += len(matches)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        matches = regex.findall(run.text)
                        if matches:
                            run.text = regex.sub(replace_text, run.text)
                            count += len(matches)

    return count


def append_paragraph(doc, text, style="Normal", **kwargs):
    """Append a paragraph to the end of the document."""
    from docx.shared import Pt, RGBColor

    p = doc.add_paragraph()
    try:
        p.style = doc.styles[style]
    except KeyError:
        pass

    run = p.add_run(text)

    if kwargs.get('bold'):
        run.bold = True
    if kwargs.get('italic'):
        run.italic = True
    if kwargs.get('font_name'):
        run.font.name = kwargs['font_name']
    if kwargs.get('font_size'):
        run.font.size = Pt(kwargs['font_size'])

    return True


def insert_paragraph_after(doc, after_text, text, style="Normal"):
    """Insert a paragraph after the paragraph containing specific text."""
    from docx.oxml.ns import qn
    from lxml import etree

    for i, para in enumerate(doc.paragraphs):
        if after_text in para.text:
            new_p = doc.add_paragraph(text)
            try:
                new_p.style = doc.styles[style]
            except KeyError:
                pass

            # Move the new paragraph after the target
            para._element.addnext(new_p._element)
            return True
    return False


def delete_paragraph(doc, containing_text):
    """Delete paragraph(s) containing specific text."""
    deleted = 0
    for para in doc.paragraphs:
        if containing_text in para.text:
            p = para._element
            p.getparent().remove(p)
            deleted += 1
    return deleted


def append_table(doc, headers, rows, style="Table Grid"):
    """Append a table to the document."""
    num_cols = len(headers)
    table = doc.add_table(rows=1, cols=num_cols)
    try:
        table.style = style
    except KeyError:
        table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = str(header)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            if i < num_cols:
                row_cells[i].text = str(val)

    return True


def insert_image(doc, image_path, width_inches=4.0, after_text=None):
    """Insert an image into the document."""
    from docx.shared import Inches

    if not os.path.exists(image_path):
        return False

    if after_text:
        for para in doc.paragraphs:
            if after_text in para.text:
                # Add image paragraph after target
                new_p = doc.add_paragraph()
                run = new_p.add_run()
                run.add_picture(image_path, width=Inches(width_inches))
                para._element.addnext(new_p._element)
                return True
        return False
    else:
        doc.add_picture(image_path, width=Inches(width_inches))
        return True


def apply_style(doc, paragraph_containing, style_name):
    """Apply a style to paragraph(s) containing specific text."""
    applied = 0
    for para in doc.paragraphs:
        if paragraph_containing in para.text:
            try:
                para.style = doc.styles[style_name]
                applied += 1
            except KeyError:
                pass
    return applied


def set_font(doc, paragraph_containing, **kwargs):
    """Set font properties on paragraph(s) containing specific text."""
    from docx.shared import Pt, RGBColor

    modified = 0
    for para in doc.paragraphs:
        if paragraph_containing in para.text:
            for run in para.runs:
                if kwargs.get('font_name'):
                    run.font.name = kwargs['font_name']
                if kwargs.get('font_size'):
                    run.font.size = Pt(kwargs['font_size'])
                if kwargs.get('bold') is not None:
                    run.bold = kwargs['bold']
                if kwargs.get('italic') is not None:
                    run.italic = kwargs['italic']
                if kwargs.get('color'):
                    hex_c = kwargs['color'].lstrip('#')
                    run.font.color.rgb = RGBColor(
                        int(hex_c[0:2], 16), int(hex_c[2:4], 16), int(hex_c[4:6], 16)
                    )
            modified += 1
    return modified


def update_header(doc, text):
    """Update header text in all sections."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        if header.paragraphs:
            header.paragraphs[0].text = text
        else:
            header.add_paragraph(text)
    return True


def update_footer(doc, text):
    """Update footer text in all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            footer.paragraphs[0].text = text
        else:
            footer.add_paragraph(text)
    return True


def update_metadata(doc, **kwargs):
    """Update document core properties."""
    props = doc.core_properties
    if 'title' in kwargs:
        props.title = kwargs['title']
    if 'author' in kwargs:
        props.author = kwargs['author']
    if 'subject' in kwargs:
        props.subject = kwargs['subject']
    if 'keywords' in kwargs:
        props.keywords = kwargs['keywords']
    if 'category' in kwargs:
        props.category = kwargs['category']
    if 'comments' in kwargs:
        props.comments = kwargs['comments']
    return True


def process_operations(input_path, output_path, operations):
    """Process a list of edit operations on a document."""
    ensure_dependencies()
    from docx import Document

    doc = Document(input_path)
    results = []

    for op in operations:
        action = op.get('action', '')
        result = {"action": action, "status": "success"}

        try:
            if action == 'find_replace':
                count = find_replace(doc, op['find'], op['replace'], op.get('match_case', True))
                result["replacements"] = count
                if count == 0:
                    result["status"] = "no_matches"
                    result["message"] = f"No matches found for '{op['find']}'"

            elif action == 'find_replace_regex':
                count = find_replace_regex(doc, op['pattern'], op['replace'])
                result["replacements"] = count

            elif action == 'append_paragraph':
                append_paragraph(doc, op['text'], op.get('style', 'Normal'),
                               bold=op.get('bold'), italic=op.get('italic'),
                               font_name=op.get('font_name'), font_size=op.get('font_size'))

            elif action == 'insert_paragraph':
                success = insert_paragraph_after(doc, op['after_text'], op['text'], op.get('style', 'Normal'))
                if not success:
                    result["status"] = "warning"
                    result["message"] = f"Could not find text: '{op['after_text']}'"

            elif action == 'delete_paragraph':
                deleted = delete_paragraph(doc, op['containing_text'])
                result["deleted"] = deleted

            elif action == 'append_table':
                append_table(doc, op['headers'], op['rows'], op.get('style', 'Table Grid'))

            elif action == 'insert_image':
                success = insert_image(doc, op['image_path'], op.get('width_inches', 4.0), op.get('after_text'))
                if not success:
                    result["status"] = "error"
                    result["message"] = "Image not found or insertion point not found"

            elif action == 'apply_style':
                count = apply_style(doc, op['paragraph_containing'], op['style'])
                result["applied"] = count

            elif action == 'set_font':
                count = set_font(doc, op['paragraph_containing'],
                               font_name=op.get('font_name'),
                               font_size=op.get('font_size'),
                               bold=op.get('bold'),
                               italic=op.get('italic'),
                               color=op.get('color'))
                result["modified"] = count

            elif action == 'update_header':
                update_header(doc, op['text'])

            elif action == 'update_footer':
                update_footer(doc, op['text'])

            elif action == 'update_metadata':
                update_metadata(doc, **{k: v for k, v in op.items() if k != 'action'})

            else:
                result["status"] = "error"
                result["message"] = f"Unknown action: {action}"

        except Exception as e:
            result["status"] = "error"
            result["message"] = str(e)

        results.append(result)

    os.makedirs(os.path.dirname(output_path) or '/tmp', exist_ok=True)
    doc.save(output_path)

    return {"status": "success", "output_path": output_path, "operations": results}


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: edit_docx.py <json_operations_file_or_string>"}))
        sys.exit(1)

    input_arg = sys.argv[1]

    if os.path.isfile(input_arg):
        with open(input_arg, 'r') as f:
            spec = json.load(f)
    else:
        spec = json.loads(input_arg)

    input_path = spec.get('input_path')
    output_path = spec.get('output_path', input_path.replace('.docx', '_edited.docx') if input_path else '/tmp/edited.docx')
    operations = spec.get('operations', [])

    if not input_path or not os.path.exists(input_path):
        print(json.dumps({"error": f"Input file not found: {input_path}"}))
        sys.exit(1)

    result = process_operations(input_path, output_path, operations)
    print(json.dumps(result, indent=2))


if __name__ == '__main__':
    main()
