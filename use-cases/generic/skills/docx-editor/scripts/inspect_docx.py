#!/usr/bin/env python3
"""Inspect Word document structure, metadata, and statistics."""

import json
import sys
import os


def ensure_dependencies():
    try:
        import docx
    except ImportError:
        os.system("pip install python-docx")


def inspect_document(doc_path):
    """Get comprehensive structural information about a .docx file."""
    ensure_dependencies()
    from docx import Document
    from docx.enum.section import WD_ORIENT

    doc = Document(doc_path)

    # Metadata
    props = doc.core_properties
    metadata = {
        "title": props.title or "",
        "author": props.author or "",
        "subject": props.subject or "",
        "keywords": props.keywords or "",
        "created": str(props.created) if props.created else None,
        "modified": str(props.modified) if props.modified else None,
        "last_modified_by": props.last_modified_by or "",
        "revision": props.revision,
        "category": props.category or "",
        "comments": props.comments or "",
    }

    # Statistics
    paragraphs = doc.paragraphs
    word_count = sum(len(p.text.split()) for p in paragraphs)
    char_count = sum(len(p.text) for p in paragraphs)
    char_count_no_spaces = sum(len(p.text.replace(' ', '')) for p in paragraphs)
    non_empty_paragraphs = [p for p in paragraphs if p.text.strip()]

    statistics = {
        "paragraph_count": len(paragraphs),
        "non_empty_paragraph_count": len(non_empty_paragraphs),
        "word_count": word_count,
        "character_count": char_count,
        "character_count_no_spaces": char_count_no_spaces,
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
    }

    # Image count
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
    statistics["image_count"] = image_count

    # Styles used
    styles_used = {}
    for p in paragraphs:
        if p.style:
            name = p.style.name
            styles_used[name] = styles_used.get(name, 0) + 1

    # Table details
    tables_info = []
    for i, table in enumerate(doc.tables):
        tables_info.append({
            "index": i,
            "rows": len(table.rows),
            "columns": len(table.columns),
        })

    # Section details
    sections_info = []
    for i, section in enumerate(doc.sections):
        orient = "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait"
        sec_info = {
            "index": i,
            "orientation": orient,
            "page_width_inches": round(section.page_width.inches, 2) if section.page_width else None,
            "page_height_inches": round(section.page_height.inches, 2) if section.page_height else None,
            "top_margin_inches": round(section.top_margin.inches, 2) if section.top_margin else None,
            "bottom_margin_inches": round(section.bottom_margin.inches, 2) if section.bottom_margin else None,
            "left_margin_inches": round(section.left_margin.inches, 2) if section.left_margin else None,
            "right_margin_inches": round(section.right_margin.inches, 2) if section.right_margin else None,
        }
        sections_info.append(sec_info)

    # Headers and footers
    headers_footers = []
    for i, section in enumerate(doc.sections):
        hf = {"section": i}
        if section.header and section.header.paragraphs:
            header_text = [p.text for p in section.header.paragraphs if p.text.strip()]
            if header_text:
                hf["header"] = header_text
        if section.footer and section.footer.paragraphs:
            footer_text = [p.text for p in section.footer.paragraphs if p.text.strip()]
            if footer_text:
                hf["footer"] = footer_text
        if len(hf) > 1:
            headers_footers.append(hf)

    # Heading outline
    outline = []
    for p in paragraphs:
        if p.style and p.style.name.startswith('Heading'):
            try:
                level = int(p.style.name.split()[-1])
            except (ValueError, IndexError):
                level = 0
            outline.append({"level": level, "text": p.text})

    # File info
    file_size = os.path.getsize(doc_path)
    file_info = {
        "path": doc_path,
        "size_bytes": file_size,
        "size_readable": _format_size(file_size),
    }

    return {
        "file": file_info,
        "metadata": metadata,
        "statistics": statistics,
        "styles_used": styles_used,
        "tables": tables_info,
        "sections": sections_info,
        "headers_footers": headers_footers,
        "outline": outline,
    }


def _format_size(size_bytes):
    """Format byte size to human-readable string."""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} TB"


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: inspect_docx.py <file.docx>"}))
        sys.exit(1)

    doc_path = sys.argv[1]

    if not os.path.exists(doc_path):
        print(json.dumps({"error": f"File not found: {doc_path}"}))
        sys.exit(1)

    try:
        result = inspect_document(doc_path)
        print(json.dumps(result, indent=2, default=str))
    except Exception as e:
        print(json.dumps({"error": str(e)}))
        sys.exit(1)


if __name__ == '__main__':
    main()
