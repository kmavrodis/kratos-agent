#!/usr/bin/env python3
"""Create a Word document from a JSON content specification."""

import json
import sys
import os
import re

def ensure_dependencies():
    try:
        import docx
    except ImportError:
        os.system("pip install python-docx")
        import docx

def parse_markdown_runs(paragraph, text):
    """Parse simple markdown (bold, italic) into Word runs."""
    from docx.shared import Pt
    parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part:
            paragraph.add_run(part)


def apply_font_settings(run, item):
    """Apply font settings to a run from an item dict."""
    from docx.shared import Pt, RGBColor
    if item.get('bold'):
        run.bold = True
    if item.get('italic'):
        run.italic = True
    if item.get('underline'):
        run.underline = True
    if item.get('font_name'):
        run.font.name = item['font_name']
    if item.get('font_size'):
        run.font.size = Pt(item['font_size'])
    if item.get('color'):
        hex_color = item['color'].lstrip('#')
        run.font.color.rgb = RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16)
        )


def set_alignment(paragraph, alignment_str):
    """Set paragraph alignment from a string."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    alignments = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if alignment_str and alignment_str.lower() in alignments:
        paragraph.alignment = alignments[alignment_str.lower()]


def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    import docx.oxml.ns as ns

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.ns.qn('r:id'), r_id)

    new_run = docx.oxml.OxmlElement('w:r')
    rPr = docx.oxml.OxmlElement('w:rPr')

    c = docx.oxml.OxmlElement('w:color')
    c.set(docx.oxml.ns.qn('w:val'), '0563C1')
    rPr.append(c)

    u = docx.oxml.OxmlElement('w:u')
    u.set(docx.oxml.ns.qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def create_document(spec):
    """Create a .docx document from a content specification dict."""
    ensure_dependencies()
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn

    doc = Document()

    # Page setup
    page = spec.get('page_setup', {})
    for section in doc.sections:
        if page.get('orientation') == 'landscape':
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
        if page.get('top_margin') is not None:
            section.top_margin = Inches(page['top_margin'])
        if page.get('bottom_margin') is not None:
            section.bottom_margin = Inches(page['bottom_margin'])
        if page.get('left_margin') is not None:
            section.left_margin = Inches(page['left_margin'])
        if page.get('right_margin') is not None:
            section.right_margin = Inches(page['right_margin'])

    # Metadata
    meta = spec.get('metadata', {})
    if meta.get('title'):
        doc.core_properties.title = meta['title']
    if meta.get('author'):
        doc.core_properties.author = meta['author']
    if meta.get('subject'):
        doc.core_properties.subject = meta['subject']
    if meta.get('keywords'):
        doc.core_properties.keywords = meta['keywords']

    # Custom styles
    for style_def in spec.get('styles', {}).get('custom_styles', []):
        try:
            from docx.enum.style import WD_STYLE_TYPE
            base_style = doc.styles[style_def.get('base', 'Normal')]
            new_style = doc.styles.add_style(style_def['name'], WD_STYLE_TYPE.PARAGRAPH)
            new_style.base_style = base_style
            font = new_style.font
            if style_def.get('font_name'):
                font.name = style_def['font_name']
            if style_def.get('font_size'):
                font.size = Pt(style_def['font_size'])
            if style_def.get('bold') is not None:
                font.bold = style_def['bold']
            if style_def.get('italic') is not None:
                font.italic = style_def['italic']
            if style_def.get('color'):
                hex_c = style_def['color'].lstrip('#')
                font.color.rgb = RGBColor(
                    int(hex_c[0:2], 16), int(hex_c[2:4], 16), int(hex_c[4:6], 16)
                )
        except Exception as e:
            print(f"Warning: Could not create style '{style_def.get('name')}': {e}", file=sys.stderr)

    # Remove the default empty paragraph
    if doc.paragraphs and doc.paragraphs[0].text == '':
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)

    # Content blocks
    for item in spec.get('content', []):
        item_type = item.get('type', '')

        if item_type == 'heading':
            heading = doc.add_heading(item.get('text', ''), level=item.get('level', 1))
            if item.get('alignment'):
                set_alignment(heading, item['alignment'])

        elif item_type == 'paragraph':
            p = doc.add_paragraph()
            if item.get('style'):
                try:
                    p.style = doc.styles[item['style']]
                except KeyError:
                    pass
            if item.get('alignment'):
                set_alignment(p, item['alignment'])

            text = item.get('text', '')
            if item.get('parse_markdown'):
                parse_markdown_runs(p, text)
            else:
                run = p.add_run(text)
                apply_font_settings(run, item)

        elif item_type == 'table':
            headers = item.get('headers', [])
            rows = item.get('rows', [])
            num_cols = len(headers) if headers else (len(rows[0]) if rows else 0)

            table = doc.add_table(rows=0, cols=num_cols)
            if item.get('style'):
                try:
                    table.style = item['style']
                except KeyError:
                    table.style = 'Table Grid'
            else:
                table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            if headers:
                hdr_row = table.add_row()
                for i, header_text in enumerate(headers):
                    cell = hdr_row.cells[i]
                    cell.text = str(header_text)
                    run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(str(header_text))
                    run.bold = True

                    if item.get('header_bg_color'):
                        hex_c = item['header_bg_color'].lstrip('#')
                        shading = docx.oxml.OxmlElement('w:shd')
                        shading.set(qn('w:fill'), hex_c)
                        shading.set(qn('w:val'), 'clear')
                        cell._tc.get_or_add_tcPr().append(shading)

                    if item.get('header_font_color'):
                        hex_c = item['header_font_color'].lstrip('#')
                        run.font.color.rgb = RGBColor(
                            int(hex_c[0:2], 16), int(hex_c[2:4], 16), int(hex_c[4:6], 16)
                        )

            for row_data in rows:
                data_row = table.add_row()
                for i, cell_text in enumerate(row_data):
                    if i < num_cols:
                        data_row.cells[i].text = str(cell_text)

        elif item_type == 'image':
            if os.path.exists(item.get('path', '')):
                width = Inches(item.get('width_inches', 5.0))
                doc.add_picture(item['path'], width=width)
                if item.get('caption'):
                    caption_p = doc.add_paragraph(item['caption'])
                    caption_p.style = 'Caption' if 'Caption' in [s.name for s in doc.styles] else 'Normal'
                    set_alignment(caption_p, 'center')
            else:
                doc.add_paragraph(f"[Image not found: {item.get('path', 'unknown')}]")

        elif item_type == 'list':
            style_map = {
                'bullet': 'List Bullet',
                'numbered': 'List Number',
            }
            list_style = style_map.get(item.get('style', 'bullet'), 'List Bullet')
            for list_item in item.get('items', []):
                p = doc.add_paragraph(str(list_item), style=list_style)

        elif item_type == 'page_break':
            doc.add_page_break()

        elif item_type == 'hyperlink':
            p = doc.add_paragraph()
            add_hyperlink(p, item.get('url', '#'), item.get('text', 'Link'))

        elif item_type == 'header':
            for section in doc.sections:
                header = section.header
                header.is_linked_to_previous = False
                if header.paragraphs:
                    header.paragraphs[0].text = item.get('text', '')
                    if item.get('alignment'):
                        set_alignment(header.paragraphs[0], item['alignment'])
                else:
                    p = header.add_paragraph(item.get('text', ''))
                    if item.get('alignment'):
                        set_alignment(p, item['alignment'])

        elif item_type == 'footer':
            for section in doc.sections:
                footer = section.footer
                footer.is_linked_to_previous = False
                p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                p.text = item.get('text', '')
                if item.get('alignment'):
                    set_alignment(p, item['alignment'])
                if item.get('include_page_number'):
                    run = p.add_run()
                    fld_char_begin = docx.oxml.OxmlElement('w:fldChar')
                    fld_char_begin.set(qn('w:fldCharType'), 'begin')
                    run._r.append(fld_char_begin)

                    instr_text = docx.oxml.OxmlElement('w:instrText')
                    instr_text.set(qn('xml:space'), 'preserve')
                    instr_text.text = ' PAGE '
                    run._r.append(instr_text)

                    fld_char_end = docx.oxml.OxmlElement('w:fldChar')
                    fld_char_end.set(qn('w:fldCharType'), 'end')
                    run._r.append(fld_char_end)

    # Save
    output_path = spec.get('output_path', '/tmp/output.docx')
    os.makedirs(os.path.dirname(output_path) or '/tmp', exist_ok=True)
    doc.save(output_path)
    return {"status": "success", "output_path": output_path}


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: create_docx.py <json_spec_file_or_json_string>"}))
        sys.exit(1)

    input_arg = sys.argv[1]

    if os.path.isfile(input_arg):
        with open(input_arg, 'r') as f:
            spec = json.load(f)
    else:
        spec = json.loads(input_arg)

    result = create_document(spec)
    print(json.dumps(result, indent=2))


if __name__ == '__main__':
    main()
